"""DOCX ingestion using python-docx.

Preserves structural elements (headings, paragraphs, tables) and returns
LangChain Document objects — one per logical section — so downstream
chunking and comparison chains have meaningful boundaries to work with.
"""

from __future__ import annotations

import logging
from pathlib import Path
from typing import Any

from docx import Document as DocxDocument
from docx.oxml.ns import qn
from langchain_core.documents import Document

logger = logging.getLogger(__name__)

# Heading styles used by python-docx (covers built-in Word styles)
_HEADING_STYLES: frozenset[str] = frozenset(
    {f"Heading {i}" for i in range(1, 10)} | {"Title", "Subtitle"}
)


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------


def load_docx(file_path: str | Path) -> list[Document]:
    """Parse a DOCX file into LangChain Documents, one per section.

    Sections are delineated by heading paragraphs. Non-heading content
    before the first heading is grouped under a synthetic "Preamble" section.
    Tables within a section are serialised as Markdown-style plain text and
    appended to that section's content.

    Args:
        file_path: Path to the ``.docx`` file.

    Returns:
        A list of Document objects with ``page_content`` = section text and
        ``metadata`` containing ``source``, ``section_index``, ``heading``,
        ``heading_level``, and ``total_sections``.

    Raises:
        FileNotFoundError: If the file does not exist.
        ValueError: If the file is not a ``.docx``.
    """
    path = _validate_path(file_path, suffix=".docx")

    try:
        raw_sections = _extract_sections(path)
    except Exception as exc:  # noqa: BLE001
        logger.error("Failed to parse '%s': %s", path.name, exc)
        raise

    total = len(raw_sections)
    docs: list[Document] = []
    for idx, section in enumerate(raw_sections):
        docs.append(
            Document(
                page_content=section["content"].strip(),
                metadata={
                    "source": str(path),
                    "section_index": idx,
                    "heading": section["heading"],
                    "heading_level": section["heading_level"],
                    "total_sections": total,
                },
            )
        )

    logger.info("Loaded %d sections from '%s'.", total, path.name)
    return docs


def extract_tables_from_docx(file_path: str | Path) -> list[dict[str, Any]]:
    """Extract all tables from a DOCX as structured dicts.

    Args:
        file_path: Path to the ``.docx`` file.

    Returns:
        List of dicts: ``{"table_index": int, "data": list[list[str]]}``
    """
    path = _validate_path(file_path, suffix=".docx")
    result: list[dict[str, Any]] = []

    try:
        doc = DocxDocument(str(path))
        for tbl_idx, table in enumerate(doc.tables):
            rows = []
            for row in table.rows:
                rows.append([cell.text.strip() for cell in row.cells])
            result.append({"table_index": tbl_idx, "data": rows})
    except Exception as exc:  # noqa: BLE001
        logger.error("Failed to extract tables from '%s': %s", path.name, exc)

    return result


# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------


def _validate_path(file_path: str | Path, suffix: str) -> Path:
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {path}")
    if path.suffix.lower() != suffix:
        raise ValueError(f"Expected a '{suffix}' file, got: {path.suffix}")
    return path


def _extract_sections(path: Path) -> list[dict[str, Any]]:
    """Walk the document body and group content into sections by heading."""
    doc = DocxDocument(str(path))

    # We iterate over the raw XML children of the body so that tables and
    # paragraphs are visited in document order (doc.paragraphs skips tables).
    body = doc.element.body

    sections: list[dict[str, Any]] = []
    current: dict[str, Any] = {
        "heading": "Preamble",
        "heading_level": 0,
        "content": "",
    }

    for child in body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

        if tag == "p":  # paragraph
            para_text = "".join(run.text for run in child.iter(qn("w:t")))
            style_name = _get_style_name(child)

            if style_name in _HEADING_STYLES:
                # Flush previous section (if it has any content)
                if current["content"].strip() or current["heading"] != "Preamble":
                    sections.append(current)
                level = _heading_level(style_name)
                current = {
                    "heading": para_text.strip(),
                    "heading_level": level,
                    "content": "",
                }
            else:
                if para_text.strip():
                    current["content"] += para_text + "\n"

        elif tag == "tbl":  # table
            table_text = _serialise_table_xml(child)
            if table_text:
                current["content"] += table_text + "\n"

    # Flush the last section
    sections.append(current)

    return sections


def _get_style_name(para_element: Any) -> str:
    """Extract the style name from a paragraph XML element."""
    try:
        pPr = para_element.find(qn("w:pPr"))
        if pPr is not None:
            pStyle = pPr.find(qn("w:pStyle"))
            if pStyle is not None:
                val = pStyle.get(qn("w:val"), "")
                # Normalise e.g. "Heading1" → "Heading 1"
                return _normalise_style(val)
    except Exception:  # noqa: BLE001
        pass
    return ""


def _normalise_style(raw: str) -> str:
    """Convert compact style IDs like 'Heading1' to display names like 'Heading 1'."""
    for i in range(9, 0, -1):
        raw = raw.replace(f"Heading{i}", f"Heading {i}")
    return raw


def _heading_level(style_name: str) -> int:
    """Return the numeric heading level (1-9) or 0 for Title/Subtitle."""
    if "Title" in style_name or "Subtitle" in style_name:
        return 0
    for i in range(1, 10):
        if str(i) in style_name:
            return i
    return 0


def _serialise_table_xml(tbl_element: Any) -> str:
    """Convert a table XML element to a simple pipe-delimited text representation."""
    rows: list[str] = []
    for tr in tbl_element.iter(qn("w:tr")):
        cells: list[str] = []
        for tc in tr.iter(qn("w:tc")):
            cell_text = "".join(t.text or "" for t in tc.iter(qn("w:t")))
            cells.append(cell_text.strip())
        if cells:
            rows.append(" | ".join(cells))
    return "\n".join(rows)
