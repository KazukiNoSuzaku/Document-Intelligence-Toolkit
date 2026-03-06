"""End-to-end pipeline smoke test.

Exercises the full path: parse → chunk → summarise → extract → compare,
with all LLM calls mocked so no API key or network access is required.

Run with:  pytest tests/test_smoke.py -v
"""

from __future__ import annotations

import tempfile
from pathlib import Path
from unittest.mock import MagicMock, patch

import pytest
from docx import Document as DocxDocument
from langchain_core.documents import Document

from src.document_parsers.docx_parser import load_docx
from src.intelligence.comparator import ComparisonResult, compare_documents
from src.intelligence.extractor import DocumentExtraction, extract_structured_data
from src.intelligence.summarizer import summarize_documents
from src.utils.chunker import chunk_documents
from src.utils.token_counter import count_tokens_for_documents


# ---------------------------------------------------------------------------
# Synthetic fixtures
# ---------------------------------------------------------------------------


def _make_docx(tmp_path: Path, name: str, body: str, heading: str = "Introduction") -> Path:
    """Write a minimal DOCX with one heading and one paragraph."""
    doc = DocxDocument()
    doc.add_heading(heading, level=1)
    doc.add_paragraph(body)
    path = tmp_path / name
    doc.save(str(path))
    return path


CONTRACT_A = (
    "This Software Services Agreement ('Agreement') is entered into as of 2024-01-15 "
    "by and between Acme Corporation ('Client') and Globex Solutions Ltd ('Provider'). "
    "The Provider agrees to deliver a document intelligence platform within 90 days. "
    "Payment of $150,000 is due within 30 days of milestone acceptance. "
    "Either party may terminate with 30 days written notice."
)

CONTRACT_B = (
    "This Software Services Agreement ('Agreement') is entered into as of 2024-03-01 "
    "by and between Acme Corporation ('Client') and Globex Solutions Ltd ('Provider'). "
    "The Provider agrees to deliver a document intelligence platform within 60 days. "
    "Payment of $175,000 is due within 14 days of milestone acceptance. "
    "Either party may terminate with 14 days written notice. "
    "Late payments accrue interest at 1.5% per month."
)

MOCK_EXTRACTION = DocumentExtraction(
    title="Software Services Agreement",
    document_type="contract",
    parties=["Acme Corporation", "Globex Solutions Ltd"],
    dates=["2024-01-15"],
    key_topics=["software development", "payment terms", "termination"],
    key_clauses=["payment within 30 days", "90-day delivery", "30-day termination notice"],
    summary="A fixed-price software services contract between Acme and Globex.",
)


# ---------------------------------------------------------------------------
# Smoke test
# ---------------------------------------------------------------------------


class TestFullPipeline:
    @pytest.fixture()
    def docx_a(self, tmp_path: Path) -> Path:
        return _make_docx(tmp_path, "contract_v1.docx", CONTRACT_A)

    @pytest.fixture()
    def docx_b(self, tmp_path: Path) -> Path:
        return _make_docx(tmp_path, "contract_v2.docx", CONTRACT_B, heading="Revised Agreement")

    def test_parse_yields_documents(self, docx_a: Path) -> None:
        docs = load_docx(docx_a)
        assert len(docs) >= 1
        assert all(isinstance(d, Document) for d in docs)

    def test_chunk_after_parse(self, docx_a: Path) -> None:
        docs = load_docx(docx_a)
        chunks = chunk_documents(docs, chunk_size=200, chunk_overlap=20)
        assert len(chunks) >= 1
        for chunk in chunks:
            assert "chunk_index" in chunk.metadata
            assert "chunk_total" in chunk.metadata

    def test_token_count_after_parse(self, docx_a: Path) -> None:
        docs = load_docx(docx_a)
        total = count_tokens_for_documents(docs)
        assert total > 0

    @patch("src.intelligence.summarizer.get_llm")
    def test_summarise(self, mock_get_llm: MagicMock, docx_a: Path) -> None:
        mock_chain = MagicMock()
        mock_chain.invoke.return_value = "A software development contract summary."
        with patch("src.intelligence.summarizer._DIRECT_PROMPT") as mock_prompt:
            mock_prompt.__or__ = MagicMock(return_value=mock_chain)
            docs = load_docx(docx_a)
            result = summarize_documents(docs, style="concise", map_reduce_threshold=999_999)
        assert isinstance(result, str)

    @patch("src.intelligence.extractor.get_llm")
    def test_extract(self, mock_get_llm: MagicMock, docx_a: Path) -> None:
        mock_llm = MagicMock()
        mock_structured = MagicMock()
        mock_chain = MagicMock()
        mock_chain.invoke.return_value = MOCK_EXTRACTION
        mock_structured.__or__ = MagicMock(return_value=mock_chain)
        mock_llm.with_structured_output.return_value = mock_structured
        mock_get_llm.return_value = mock_llm

        with patch("src.intelligence.extractor._EXTRACTION_PROMPT") as mock_prompt:
            mock_prompt.__or__ = MagicMock(return_value=mock_chain)
            docs = load_docx(docx_a)
            result = extract_structured_data(docs)

        assert isinstance(result, DocumentExtraction)
        assert result.document_type == "contract"
        assert "Acme Corporation" in result.parties

    @patch("src.intelligence.comparator.get_llm")
    def test_compare(self, mock_get_llm: MagicMock, docx_a: Path, docx_b: Path) -> None:
        mock_chain = MagicMock()
        mock_chain.invoke.return_value = (
            "There are significant changes: delivery shortened from 90 to 60 days, "
            "payment increased to $175,000, and a late-payment clause was added."
        )
        with patch("src.intelligence.comparator._SEMANTIC_DIFF_PROMPT") as mock_prompt:
            mock_prompt.__or__ = MagicMock(return_value=mock_chain)
            docs_a = load_docx(docx_a)
            docs_b = load_docx(docx_b)
            result = compare_documents(docs_a, docs_b)

        assert isinstance(result, ComparisonResult)
        assert result.exact_diff.total_changes > 0
        assert result.similarity_assessment == "significant changes"
        assert isinstance(result.semantic_summary, str)

    @patch("src.intelligence.comparator.get_llm")
    @patch("src.intelligence.extractor.get_llm")
    @patch("src.intelligence.summarizer.get_llm")
    def test_full_pipeline_single_document(
        self,
        mock_sum_llm: MagicMock,
        mock_ext_llm: MagicMock,
        mock_cmp_llm: MagicMock,
        docx_a: Path,
        docx_b: Path,
    ) -> None:
        """Full parse → chunk → summarise → extract → compare pipeline."""
        # --- Parse ---
        docs_a = load_docx(docx_a)
        docs_b = load_docx(docx_b)
        assert docs_a and docs_b

        # --- Chunk ---
        chunks = chunk_documents(docs_a, chunk_size=300)
        assert chunks

        # --- Summarise ---
        sum_chain = MagicMock()
        sum_chain.invoke.return_value = "Contract summary."
        with patch("src.intelligence.summarizer._DIRECT_PROMPT") as sp:
            sp.__or__ = MagicMock(return_value=sum_chain)
            summary = summarize_documents(docs_a, map_reduce_threshold=999_999)
        assert summary

        # --- Extract ---
        ext_llm = MagicMock()
        ext_structured = MagicMock()
        ext_chain = MagicMock()
        ext_chain.invoke.return_value = MOCK_EXTRACTION
        ext_structured.__or__ = MagicMock(return_value=ext_chain)
        ext_llm.with_structured_output.return_value = ext_structured
        mock_ext_llm.return_value = ext_llm
        with patch("src.intelligence.extractor._EXTRACTION_PROMPT") as ep:
            ep.__or__ = MagicMock(return_value=ext_chain)
            extraction = extract_structured_data(docs_a)
        assert extraction.title

        # --- Compare ---
        cmp_chain = MagicMock()
        cmp_chain.invoke.return_value = "significant changes to payment and timeline."
        with patch("src.intelligence.comparator._SEMANTIC_DIFF_PROMPT") as cp:
            cp.__or__ = MagicMock(return_value=cmp_chain)
            comparison = compare_documents(docs_a, docs_b)
        assert comparison.exact_diff.total_changes >= 0
        assert comparison.semantic_summary
