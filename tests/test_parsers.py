"""Unit tests for PDF and DOCX parsers.

Tests use only in-memory / on-disk synthetic fixtures — no real LLM calls.
Run with:  pytest tests/test_parsers.py -v
"""

from __future__ import annotations

import io
import struct
import textwrap
import zlib
from pathlib import Path

import pytest
from docx import Document as DocxDocument
from langchain_core.documents import Document

from src.document_parsers.pdf_parser import (
    _clean_metadata,
    _validate_path,
    extract_tables_from_pdf,
    load_pdf,
)
from src.document_parsers.docx_parser import (
    extract_tables_from_docx,
    load_docx,
)


# ---------------------------------------------------------------------------
# Fixtures — synthetic files written to a tmp directory
# ---------------------------------------------------------------------------


@pytest.fixture()
def sample_docx(tmp_path: Path) -> Path:
    """Create a minimal DOCX with headings, paragraphs, and a table."""
    doc = DocxDocument()
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph("This is the introduction section.")
    doc.add_heading("Scope of Work", level=2)
    doc.add_paragraph("The scope covers document intelligence features.")
    # Add a simple 2x2 table
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Feature"
    table.cell(0, 1).text = "Status"
    table.cell(1, 0).text = "PDF Parsing"
    table.cell(1, 1).text = "Done"
    file_path = tmp_path / "sample.docx"
    doc.save(str(file_path))
    return file_path


@pytest.fixture()
def sample_pdf(tmp_path: Path) -> Path:
    """Create a minimal but valid single-page PDF using raw PDF syntax."""
    content = b"Hello from a test PDF. This is page one."
    pdf_bytes = _build_minimal_pdf(content)
    file_path = tmp_path / "sample.pdf"
    file_path.write_bytes(pdf_bytes)
    return file_path


@pytest.fixture()
def nonexistent_path(tmp_path: Path) -> Path:
    return tmp_path / "ghost.pdf"


@pytest.fixture()
def wrong_extension(tmp_path: Path) -> Path:
    p = tmp_path / "document.txt"
    p.write_text("not a pdf")
    return p


# ---------------------------------------------------------------------------
# Helper — minimal valid single-page PDF generator
# ---------------------------------------------------------------------------


def _build_minimal_pdf(page_text: bytes) -> bytes:
    """Build the smallest possible valid PDF that pdfplumber can open."""
    # Encode text in a content stream
    stream_content = b"BT /F1 12 Tf 72 720 Td (" + page_text + b") Tj ET"
    stream_len = len(stream_content)

    lines = [
        b"%PDF-1.4",
        b"1 0 obj << /Type /Catalog /Pages 2 0 R >> endobj",
        b"2 0 obj << /Type /Pages /Kids [3 0 R] /Count 1 >> endobj",
        b"3 0 obj << /Type /Page /Parent 2 0 R"
        b" /MediaBox [0 0 612 792]"
        b" /Contents 4 0 R"
        b" /Resources << /Font << /F1 5 0 R >> >> >> endobj",
        b"4 0 obj << /Length " + str(stream_len).encode() + b" >>",
        b"stream",
        stream_content,
        b"endstream endobj",
        b"5 0 obj << /Type /Font /Subtype /Type1 /BaseFont /Helvetica >> endobj",
    ]

    body = b"\n".join(lines) + b"\n"

    # Cross-reference table
    xref_offset = len(body)
    # We won't build a proper xref (that requires byte offsets per object),
    # so use a linearised-enough structure that pdfplumber tolerates.
    trailer = (
        b"xref\n0 6\n"
        b"0000000000 65535 f \n"
        + b"\n".join(b"0000000000 00000 n " for _ in range(5))
        + b"\n"
        b"trailer << /Size 6 /Root 1 0 R >>\n"
        b"startxref\n" + str(xref_offset).encode() + b"\n%%EOF"
    )

    return body + trailer


# ---------------------------------------------------------------------------
# _validate_path tests
# ---------------------------------------------------------------------------


class TestValidatePath:
    def test_raises_for_missing_file(self, nonexistent_path: Path) -> None:
        with pytest.raises(FileNotFoundError, match="File not found"):
            _validate_path(nonexistent_path, ".pdf")

    def test_raises_for_wrong_extension(self, wrong_extension: Path) -> None:
        with pytest.raises(ValueError, match="Expected a '.pdf' file"):
            _validate_path(wrong_extension, ".pdf")

    def test_returns_path_object_for_valid_file(self, sample_pdf: Path) -> None:
        result = _validate_path(sample_pdf, ".pdf")
        assert isinstance(result, Path)
        assert result == sample_pdf


# ---------------------------------------------------------------------------
# _clean_metadata tests
# ---------------------------------------------------------------------------


class TestCleanMetadata:
    def test_stringifies_values(self) -> None:
        raw = {"/Author": "Alice", "/Pages": 42, "/Empty": None}
        cleaned = _clean_metadata(raw)
        assert cleaned["Author"] == "Alice"
        assert cleaned["Pages"] == "42"
        assert cleaned["Empty"] == ""

    def test_strips_leading_slash_from_keys(self) -> None:
        raw = {"/Title": "Doc"}
        cleaned = _clean_metadata(raw)
        assert "Title" in cleaned
        assert "/Title" not in cleaned

    def test_handles_empty_dict(self) -> None:
        assert _clean_metadata({}) == {}


# ---------------------------------------------------------------------------
# load_pdf tests (uses fallback to pypdf for the minimal synthetic PDF)
# ---------------------------------------------------------------------------


class TestLoadPdf:
    def test_returns_list_of_documents(self, sample_pdf: Path) -> None:
        docs = load_pdf(sample_pdf)
        assert isinstance(docs, list)
        assert len(docs) >= 1
        assert all(isinstance(d, Document) for d in docs)

    def test_metadata_contains_source(self, sample_pdf: Path) -> None:
        docs = load_pdf(sample_pdf)
        for doc in docs:
            assert "source" in doc.metadata
            assert str(sample_pdf) == doc.metadata["source"]

    def test_metadata_contains_page(self, sample_pdf: Path) -> None:
        docs = load_pdf(sample_pdf)
        for doc in docs:
            assert "page" in doc.metadata
            assert isinstance(doc.metadata["page"], int)

    def test_raises_for_missing_file(self, tmp_path: Path) -> None:
        with pytest.raises(FileNotFoundError):
            load_pdf(tmp_path / "missing.pdf")

    def test_raises_for_wrong_extension(self, tmp_path: Path) -> None:
        p = tmp_path / "file.docx"
        p.write_text("nope")
        with pytest.raises(ValueError):
            load_pdf(p)

    def test_accepts_string_path(self, sample_pdf: Path) -> None:
        docs = load_pdf(str(sample_pdf))
        assert len(docs) >= 1


# ---------------------------------------------------------------------------
# load_docx tests
# ---------------------------------------------------------------------------


class TestLoadDocx:
    def test_returns_list_of_documents(self, sample_docx: Path) -> None:
        docs = load_docx(sample_docx)
        assert isinstance(docs, list)
        assert len(docs) >= 1
        assert all(isinstance(d, Document) for d in docs)

    def test_detects_headings_as_section_boundaries(self, sample_docx: Path) -> None:
        docs = load_docx(sample_docx)
        headings = [d.metadata["heading"] for d in docs]
        assert "Introduction" in headings
        assert "Scope of Work" in headings

    def test_metadata_contains_required_fields(self, sample_docx: Path) -> None:
        docs = load_docx(sample_docx)
        required = {"source", "section_index", "heading", "heading_level", "total_sections"}
        for doc in docs:
            assert required <= doc.metadata.keys(), (
                f"Missing keys in metadata: {required - doc.metadata.keys()}"
            )

    def test_section_content_is_non_empty_for_heading_sections(
        self, sample_docx: Path
    ) -> None:
        docs = load_docx(sample_docx)
        intro = next(d for d in docs if d.metadata["heading"] == "Introduction")
        assert "introduction" in intro.page_content.lower()

    def test_table_content_included_in_section(self, sample_docx: Path) -> None:
        docs = load_docx(sample_docx)
        # The table is after "Scope of Work" heading, so it should be in that section
        full_text = " ".join(d.page_content for d in docs)
        assert "PDF Parsing" in full_text

    def test_raises_for_missing_file(self, tmp_path: Path) -> None:
        with pytest.raises(FileNotFoundError):
            load_docx(tmp_path / "ghost.docx")

    def test_raises_for_wrong_extension(self, tmp_path: Path) -> None:
        p = tmp_path / "file.pdf"
        p.write_bytes(b"%PDF-1.4")
        with pytest.raises(ValueError):
            load_docx(p)

    def test_accepts_string_path(self, sample_docx: Path) -> None:
        docs = load_docx(str(sample_docx))
        assert len(docs) >= 1


# ---------------------------------------------------------------------------
# extract_tables_from_docx tests
# ---------------------------------------------------------------------------


class TestExtractTablesFromDocx:
    def test_extracts_correct_number_of_tables(self, sample_docx: Path) -> None:
        tables = extract_tables_from_docx(sample_docx)
        assert len(tables) == 1

    def test_table_structure(self, sample_docx: Path) -> None:
        tables = extract_tables_from_docx(sample_docx)
        tbl = tables[0]
        assert "table_index" in tbl
        assert "data" in tbl
        assert tbl["table_index"] == 0

    def test_table_data_content(self, sample_docx: Path) -> None:
        tables = extract_tables_from_docx(sample_docx)
        flat = [cell for row in tables[0]["data"] for cell in row]
        assert "Feature" in flat
        assert "PDF Parsing" in flat

    def test_returns_empty_list_for_docx_without_tables(self, tmp_path: Path) -> None:
        doc = DocxDocument()
        doc.add_paragraph("No tables here.")
        path = tmp_path / "notables.docx"
        doc.save(str(path))
        tables = extract_tables_from_docx(path)
        assert tables == []


# ---------------------------------------------------------------------------
# extract_tables_from_pdf tests
# ---------------------------------------------------------------------------


class TestExtractTablesFromPdf:
    def test_returns_list(self, sample_pdf: Path) -> None:
        tables = extract_tables_from_pdf(sample_pdf)
        assert isinstance(tables, list)

    def test_table_dict_structure(self, sample_pdf: Path) -> None:
        """Each table entry must have page, table_index, and data keys."""
        tables = extract_tables_from_pdf(sample_pdf)
        for tbl in tables:
            assert "page" in tbl
            assert "table_index" in tbl
            assert "data" in tbl

    def test_raises_for_missing_file(self, tmp_path: Path) -> None:
        with pytest.raises(FileNotFoundError):
            extract_tables_from_pdf(tmp_path / "missing.pdf")

    def test_raises_for_wrong_extension(self, tmp_path: Path) -> None:
        p = tmp_path / "file.txt"
        p.write_text("not a pdf")
        with pytest.raises(ValueError):
            extract_tables_from_pdf(p)
